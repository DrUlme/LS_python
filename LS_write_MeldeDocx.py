#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Aug 29 19:24:40 2021

pip install python-docx

@author: ulf
"""

from docx import Document
from docx.shared import Inches
import sqlite3
import numpy as np

# lade globale Parameter
import LSglobal

#________________________________________________________
# Verbindung zur Datenbank erzeugen
connection = sqlite3.connect( LSglobal.SQLiteFile )
# Datensatzcursor erzeugen
cursor = connection.cursor()

#________________________________________________________
document = Document()

# Kopfzeile:
# section = document.sections[0]   
# header = section.header   
# header_para = header.paragraphs[0]   
# header_para.text = "23.10.2021\tHerbst-Langsteckentest Erlangen\tBild?"
# NOT: header_para.add_picture('RVE_BRV_Flag.png', width=Inches(0.75))

header = document.sections[0].header
paragraph = header.paragraphs[0]


text_run = paragraph.add_run()
text_run.text = "BRV Langsteckentest Erlangen " + LSglobal.Zeit + " '" + str(LSglobal.Jahr-2000) +  '\t' # For center align of text
text_run.style = "Heading 2 Char"
logo_run = paragraph.add_run()
logo_run.add_picture("RVE_BRV_Flag.png", width=Inches(1.25))

# Dokument Überschrift
document.add_heading('Ausschreibung', 0)
# Langstreckentest des Bayerischen Ruderverbandes auf dem Main-Donau-Kanal in Erlangen
p = document.add_paragraph()
p.add_run('Langstreckentest des Bayerischen Ruderverbandes \nauf dem Main-Donau-Kanal in Erlangen\n\n').bold = True
p.add_run('Ausrichter: \t\t').bold = True
p.add_run('Ruderverein Erlangen e.V. 1911 ')
p.add_run(' (RVE)\n').italic = True

p.add_run('Termin:\t\t\t').bold = True
p.add_run('Samstag, den ' +  LSglobal.Datum + " " + str(LSglobal.Jahr) + '\n')
          
p.add_run('Veranstaltungsort:\t').bold = True
p.add_run('RVE, Habichtstr.12, 91056 Erlangen\n')
p.add_run('Meldeschluss:\t\t').bold = True
p.add_run('Montag, den ').bold = True
p.add_run( LSglobal.Meldeschluss + " " + str(LSglobal.Jahr) + '  ')
p.add_run('22:00 Uhr\n').italic = True
p.add_run('Meldungen an:\t').bold = True
p.add_run('langstrecke@ruderverein-erlangen.de\n\n').italic = True
p.add_run('Bitte bei der Meldung das aktuelle elektronische Meldeformular von der Homepage des BRV nutzen! ').bold = True
p.add_run('Es muss das Rennen, Verein / Renngemeinschaft, Geburtsjahr und evtl. Leichtgewicht angegeben werden.')


# p.add_run('italic.').italic = True

document.add_heading('Rennen', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
# document.add_paragraph(    'first item in unordered list', style='List Bullet')
# document.add_paragraph(    'first item in ordered list', style='List Number')
# document.add_picture('RVE_BRV_Flag.png', width=Inches(2.75))
list1 = (0, '                ', ' ', '               ' )
# conversion
array1 = np.asarray(list1)
array2 = array1

listSize = 0


# Ausgabe des Ergebnisses
# Sammeln für 6000 m
sql = "SELECT * FROM rennen WHERE strecke='6000 m'"
cursor.execute(sql)
iD = -1
for dsatz in cursor:
   Rennen = dsatz[0]
   # R_Name = dsatz[1]
   iD = iD + 1
   if(iD > listSize):
       array1=np.vstack((array1, array2))
       listSize = listSize + 1
   if(listSize == 0):
      array1[0] = str( Rennen )
      array1[1] = dsatz[3]
   else:    
       array1[iD][0] = str( Rennen )
       array1[iD][1] = dsatz[3]
   
# Sammeln für 3000 m
sql = "SELECT * FROM rennen WHERE strecke='3000 m'"
cursor.execute(sql)
iD = -1
for dsatz in cursor:
   Rennen = dsatz[0]
   R_Name = dsatz[3]
   iD = iD + 1
   if(iD > listSize):
       array1=np.vstack((array1, array2))
       listSize = listSize + 1
   if(listSize == 0):
       array1[iD][2] = dsatz[0]
       array1[iD][3] = dsatz[3]
   else:
       array1[iD][2] = dsatz[0]
       array1[iD][3] = dsatz[3]
   

#document.add_paragraph()

table = document.add_table(rows=1, cols=5)
table.autofit = False 
table.allow_autofit = False
table.columns[0].width = Inches(0.6)
table.columns[1].width = Inches(2.5)
table.columns[2].width = Inches(0.4)
table.columns[3].width = Inches(0.6)
table.columns[4].width = Inches(2.5)

# deprecated: table.style = 'TableGrid'
# without error: table.style = document.styles['TableGrid']
table.style = 'LightShading-Accent1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nr'
hdr_cells[1].text = '6000 m'
hdr_cells[2].text = ''
hdr_cells[3].text = 'Nr'
hdr_cells[4].text = '3000 m'
for Nr, L6000m, Nr2, L3000m in array1:
    row_cells = table.add_row().cells
    row_cells[0].text = str(Nr)
    row_cells[1].text = L6000m
    row_cells[3].text = str(Nr2)
    row_cells[4].text = L3000m
#________________________________________________________
p = document.add_paragraph()
p.add_run('\nMit seiner Meldung stimmt jeder Teilnehmer zu, dass die vom Veranstalter angefertigten Foto-, Ton- und Filmaufnahmen im Zusammenhang mit der Veranstaltung veröffentlicht werden.').bold = True
#
#
document.add_page_break()
document.add_heading('Hinweise:', level=1)
pt = document.add_paragraph('Für Jungen und Mädchen sind die Rennen über 3000m im Einer und Doppelzweier vorgesehen.', style='List Number')
pt = document.add_paragraph('Die Durchführung der Langstrecke in einem Block hat sich bewährt. Es besteht danach die Möglichkeit Großboote zu testen, denn wie bekannt hält sich die Schifffahrt in Grenzen.', style='List Number')
pt = document.add_paragraph('Schifffahrtsperre: 11:00 Uhr bis 14:00 Uhr\nStartzeit: 11:00 Uhr bis ca. 13:40 Uhr ', style='List Number')
pt = document.add_paragraph('Der Betriebsweg ist zu jeder Zeit von Booten, Anhängern und Fahrzeugen freizuhalten. Die Fahrzeuge sind so abzustellen, dass eine ungehinderte Durchfahrt möglich ist. Die Kraftfahrzeuge sind auf dem Parkplatz abzustellen.', style='List Number')
pt = document.add_paragraph('', style='List Number')
pt.add_run('Die Fahrordnung ist einzuhalten!\n').bold = True
pt.add_run('Weisen Sie Ihrer Ruderer darauf hin, dass in den Wettkampfpausen die Schifffahrt nicht behindert werden darf. Ein Queren des Kanals ist nur bei ausreichendem Abstand (ca. 300m) vor der Großschifffahrt möglich. ')
pt.add_run('Wir müssen auf die Großschifffahrt Rücksicht nehmen und nicht umgekehrt. \nDen Ordnern ist Folge zu leisten.\n').bold = True
pt.add_run('Bei der Fahrt zum Start und nach den Rennen nicht nebeneinander fahren, bzw. zügig überholen und das Überholen ermöglichen.\n').bold = True
pt.add_run('Die Kanalmitte immer frei halten!').bold = True
pt = document.add_paragraph('Eine Doppelnutzung von Booten ist begrenzt möglich! Bitte mit der Meldung angeben! Bei Doppelnutzung sollten Nichtkadermitglieder im Rennen Doppelnutzung gemeldet werden.', style='List Number')
#
#p = document.add_paragraph('X')
#super_text = p.add_run('h')
#super_text.font.superscript = True
#sub_text = p.add_run('low')
#sub_text.font.subscript = True
#
document.add_heading('Besondere Bestimmungen:', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
#document.value="restart"
# style = document.styles['List Number']
pt = document.add_paragraph('Es gelten die Altersklassen für ' + str(LSglobal.RefJahr) + '\n', style='List Number')
pt.add_run('Streckenlänge: 6.000m (Start bei km 42 - Ziel bei km 48) ohne Wende\n')
pt.add_run('Streckenlänge: 3.000m (Start bei km 42 - Ziel bei km 45) ohne Wende\n')
pt.add_run('Die Boote werden im Abstand von ca. 1 Minute gestartet und haben sich 5 Minuten vor der Startzeit zur Verfügung des Starters zu halten. Änderungen werden rechtzeitig mit dem Meldeergebnis bekanntgegeben. Auch wenn es zu Beginn zu Verzögerungen kommt: Schicken Sie Ihre Ruderer rechtzeitig zum Start, denn bei Bedarf kann der Rennabstand auf 30 Sekunden verringert werden.\n')
pt.add_run('Der Start erfolgt von der Startlinie und sofort nach dem      Ausrichten der Boote.')
#
pt = document.add_paragraph('Bugnummern und Rückennummern werden kostenlos ausgegeben. Bei Verlust sind pro Start-Nr. € 10,00 zu zahlen. \n', style='List Number')
pt.add_run('Die Bugnummern werden am Steg wieder eingesammelt.').bold = True
#
pt = document.add_paragraph('Wegen Corona und dem Platz am Steg sollte diesmal erneut auf Großboote verzichtet werden, 2x/2- sind aber explizit erlaubt!', style='List Number')
#
pt = document.add_paragraph('Der Leistungstest wird nach den RWR des DRV und den Bestimmungen für die Durchführung von Jungen- und Mädchen-Wettbewerben ausgetragen soweit diese auf den Wettbewerb anwendbar sind. ', style='List Number')
pt.add_run('\nFür Kader-Angehörige des BRV und Junioren/-innen Altersklasse B, die an der Mannschaftsbildung teilnehmen wollen, ist die Teilnahme verbindlich.').bold = True
#
pt = document.add_paragraph('Alle Leichtgewichte werden gewogen. Es gelten das Gewicht laut RWR+2,5kg und die Gewichtsklassen nach den Bestimmungen für die Durchführung von Jungen-und Mädchen-Wettbewerben. ', style='List Number')
pt.add_run('Um Gruppenbildung an der Waage zu vermeiden, entfällt die 2-Std-Regel. Es kann also schon bei Ankunft gewogen werden. ')
#
pt = document.add_paragraph('Leihboote können nicht zur Verfügung gestellt werden. Die Lagerung der Boote erfolgt auf eigene Gefahr. Eine besondere Versicherung für Teilnehmer und Boote durch den Veranstalter besteht nicht.', style='List Number')
pt.add_run('Bei Unfällen und Schäden jeglicher Art haften die meldenden Teilnehmer und Vereine.')
#
pt = document.add_paragraph('Die Fahrordnung ist unbedingt einzuhalten. Für die Sicherheit der Ruderer und die Einhaltung der Fahrordnung sind neben den Ruderern auch die Trainer verantwortlich.', style='List Number')
pt.add_run('Es wird empfohlen, die unerfahrenen und unsicheren Ruderer mit Rettungswesten starten zu lassen.')
pt.add_run('Eine Teilnahme erfolgt auf eigene Gefahr. Eine Haftung durch den Veranstalter erfolgt nicht. Bei Unfällen und Schäden jeglicher Art haften die meldenden Teilnehmer und Vereine.')
pt.add_run('Den überholenden Booten ist der Überholvorgang durch ausweichen zu erleichtern.')
pt.add_run('Für Schäden wird vom Veranstalter keine Haftung übernommen.\n')
#
#
document.add_heading('Startreihenfolge:', level=1)
pt = document.add_paragraph('Die Ruderer und Ruderinnen werden entsprechend der Ergebnisse des Vorjahres und der Ergometer Werte, bzw. nach Rücksprache mit dem Landestrainer eingeordnet. \n')
pt.add_run('Die Startreihenfolge wird am 17. Oktober 2021 festgelegt.').bold = True
#
# document.paragraph_format.keep_with_next = False
document.add_heading('Meldegeld:', level=1)
pt = document.add_paragraph('Es wird ein Unkostenbeitrag von € 15,00 je Boot erhoben.\n')
pt.add_run('Bei Meldungen nach dem ' )
pt.add_run( LSglobal.DoppeltGeld ).bold = True
pt.add_run(' ist das doppelte Meldegeld zu zahlen! \n')
pt.add_run('Um- und Abmeldungen sind nur bis Freitag den ')
pt.add_run( LSglobal.AbmeldungDD ).bold = True
pt.add_run('  ' + LSglobal.AbmeldungHH + ' möglich - danach wird das Meldegeld in Rechnung gestellt. \n')
pt.add_run('Der maximale Beitrag pro Verein beträgt € 250,00 zuzüglich eventueller Nachmeldungen.\n')
pt.add_run('Die Meldeliste, Rückennummern und Bugnummern sind in einem Umschlag pro Verein im Regattabüro abholbar. ').bold = True
pt.add_run(' Die Rechnung wird nach der Regatta per Mail den Trainern zugeschickt, verlorene Bugnummern dem Verein in Rechnung gestellt. ' +
           ' Das Meldegeld wird nach der Regatta überwiesen. \n')

document.add_heading('Regelung zur Covid19 Pandemie:', level=1)
pt = document.add_paragraph('Bei der Veranstaltung sind die geltenden Regelungen zur Bekämpfung der Covid19-Pandemie zwingend einzuhalten. Erkrankte oder symptombehaftete Personen sind von der Teilnahme ausgeschlossen. Dies ' +
'gilt für alle Regelungen staatlicher oder öffentlicher, gleich ob gesetzlich, untergesetzlich oder behördlicher Einzelverfügung. Sofern diese Regelungen einer Durchführung der Veranstaltung entgegenstehen, ' +
'behält sich der Bayerische Ruderverband und RudervereinErlangen die Absage der Veranstaltung ausdrücklich vor. Für den Fall der Absage ist eine Erstattung von Kosten der Teilnehmer oder eine anderweitige Entschädigung ausgeschlossen;' +
' eine Haftung des Bayerischen Ruderverbandes und/oder des Rudervereins Erlangen besteht nicht.\n');
#
pt.add_run('Die entsprechenden Hygienevorschriften vor Ort werden mit dem Meldeergebnis versandt.').bold = True
##
pt.add_run('\n\nRV Erlangen\t\t\t\tBayerischer Ruderverband\ngez.\t\t\t\t\tgez.\nDr. Ulf Meerwald\t\t\tGerhard Walter \n2. Vorsitzender\t\t\tVizepräsident Sport')

document.save('demo.docx')
